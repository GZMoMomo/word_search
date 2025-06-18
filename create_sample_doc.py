#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建示例Word文档
用于测试文档解析API
"""

from docx import Document
from docx.shared import Inches
import os

def create_sample_document():
    """创建示例Word文档"""
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('示例文档', 0)
    
    # 添加一级标题
    doc.add_heading('第一章 介绍', level=1)
    
    # 添加正文
    doc.add_paragraph('这是第一章的介绍内容。本章将介绍文档的基本结构和内容。')
    doc.add_paragraph('文档解析API可以自动识别这些标题和内容，并提供结构化的访问方式。')
    
    # 添加二级标题
    doc.add_heading('1.1 背景', level=2)
    doc.add_paragraph('在这个背景下，我们需要一个能够解析各种文档格式的API服务。')
    doc.add_paragraph('该服务支持HTML和Word文档的解析，能够提取目录结构和章节内容。')
    
    doc.add_heading('1.2 目标', level=2)
    doc.add_paragraph('主要目标是提供一个统一的接口来访问不同格式的文档内容。')
    doc.add_paragraph('通过标准化的API，用户可以轻松获取文档的目录和特定章节的内容。')
    
    # 添加另一个一级标题
    doc.add_heading('第二章 技术实现', level=1)
    
    doc.add_paragraph('本章将详细介绍技术实现的细节。')
    
    doc.add_heading('2.1 架构设计', level=2)
    doc.add_paragraph('系统采用Flask框架构建，提供RESTful API接口。')
    doc.add_paragraph('使用python-docx库解析Word文档，使用BeautifulSoup解析HTML文档。')
    
    doc.add_heading('2.2 核心功能', level=2)
    doc.add_paragraph('核心功能包括：')
    
    # 添加列表
    features = doc.add_paragraph()
    features.add_run('• 文档类型自动识别\n')
    features.add_run('• 目录结构提取\n')
    features.add_run('• 章节内容获取\n')
    features.add_run('• 层级导航支持\n')
    
    # 添加第三个一级标题
    doc.add_heading('第三章 使用指南', level=1)
    
    doc.add_paragraph('本章提供详细的使用指南和示例。')
    
    doc.add_heading('3.1 API接口', level=2)
    doc.add_paragraph('主要提供两个API接口：')
    doc.add_paragraph('1. /api/chapters - 获取文档目录')
    doc.add_paragraph('2. /api/content - 获取章节内容')
    
    doc.add_heading('3.2 参数说明', level=2)
    doc.add_paragraph('url参数：文档的URL地址（必选）')
    doc.add_paragraph('chapter参数：章节名称（可选）')
    
    # 添加总结
    doc.add_heading('第四章 总结', level=1)
    doc.add_paragraph('本文档展示了Word文档的结构化内容。')
    doc.add_paragraph('通过使用标准的标题样式，API可以准确识别文档的层次结构。')
    doc.add_paragraph('这种结构化的文档更容易被程序解析和处理。')
    
    # 保存文档
    filename = 'sample_document.docx'
    doc.save(filename)
    
    print(f"示例文档已创建: {filename}")
    print("文档包含以下结构:")
    print("- 示例文档 (标题)")
    print("  - 第一章 介绍")
    print("    - 1.1 背景")
    print("    - 1.2 目标")
    print("  - 第二章 技术实现")
    print("    - 2.1 架构设计")
    print("    - 2.2 核心功能")
    print("  - 第三章 使用指南")
    print("    - 3.1 API接口")
    print("    - 3.2 参数说明")
    print("  - 第四章 总结")
    
    return filename

def create_simple_document():
    """创建简单的测试文档"""
    
    doc = Document()
    
    doc.add_heading('简单测试文档', 0)
    
    doc.add_heading('测试章节', level=1)
    doc.add_paragraph('这是一个测试章节的内容。')
    doc.add_paragraph('用于验证API的解析功能。')
    
    doc.add_heading('子章节', level=2)
    doc.add_paragraph('这是子章节的内容。')
    
    filename = 'simple_test.docx'
    doc.save(filename)
    
    print(f"简单测试文档已创建: {filename}")
    return filename

if __name__ == "__main__":
    print("创建示例Word文档...")
    
    # 创建详细示例文档
    sample_file = create_sample_document()
    
    print("\n" + "="*50)
    
    # 创建简单测试文档
    simple_file = create_simple_document()
    
    print(f"\n文档创建完成！")
    print(f"详细示例文档: {sample_file}")
    print(f"简单测试文档: {simple_file}")
    print("\n您可以将这些文档上传到可公开访问的URL，然后使用API进行测试。") 