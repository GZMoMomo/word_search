from flask import Flask, request, jsonify
import requests
import logging
from docx import Document
import io
import os

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

class DocumentParser:
    """文档解析器类，专门处理Word文档"""
    
    def __init__(self, url):
        self.url = url
        self.doc = None
        self._fetch_document()
    
    def _fetch_document(self):
        """获取Word文档内容"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(self.url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # 处理Word文档
            self.doc = Document(io.BytesIO(response.content))
            logger.info(f"成功获取Word文档: {self.url}")
                
        except Exception as e:
            logger.error(f"获取Word文档失败: {e}")
            raise
    
    def get_chapters(self, parent_chapter=None, return_content=False):
        """
        获取目录章节
        
        Args:
            parent_chapter (str, optional): 父级章节名称，如果为None则获取第一层级
            return_content (bool): 如果没有子章节，是否返回章节内容
            
        Returns:
            list: 章节列表，或者包含内容的字典
        """
        return self._get_word_chapters(parent_chapter, return_content)
    
    def _get_word_chapters(self, parent_chapter=None, return_content=False):
        """从Word文档中获取章节"""
        if not self.doc:
            return []
        
        chapters = []
        current_level = 0
        parent_found = parent_chapter is None
        
        for paragraph in self.doc.paragraphs:
            # 检查段落样式来判断是否为标题
            if paragraph.style.name.startswith('Heading'):
                try:
                    level = int(paragraph.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 1
                
                title = paragraph.text.strip()
                if not title:
                    continue
                
                # 如果指定了父章节，需要找到该章节的子章节
                if parent_chapter:
                    # 精确匹配父章节名称，避免误匹配
                    if title == parent_chapter:
                        parent_found = True
                        current_level = level
                        continue
                    elif parent_found:
                        # 如果已经找到父章节，检查当前章节是否是子章节
                        if level > current_level:
                            chapters.append({
                                'title': title,
                                'href': f"#heading_{len(chapters)}",
                                'level': level
                            })
                        elif level <= current_level:
                            # 遇到同级或更高级的标题，停止搜索子章节
                            break
                else:
                    # 获取第一层级章节（通常是Heading 1）
                    if level == 1:
                        chapters.append({
                            'title': title,
                            'href': f"#heading_{len(chapters)}",
                            'level': level
                        })
        
        # 如果没有找到子章节且要求返回内容，则返回父章节的内容
        if parent_chapter and not chapters and return_content:
            content = self.get_content(parent_chapter)
            if content and content != "未找到指定章节的内容":
                return {
                    'type': 'content',
                    'chapter': parent_chapter,
                    'content': content,
                    'is_leaf': True
                }
        
        if return_content:
            for chapter in chapters:
                chapter['content'] = self.get_content(chapter['title'])
        
        return chapters[:20]  # 限制返回数量
    
    def get_content(self, chapter_name=None):
        """
        获取章节内容
        
        Args:
            chapter_name (str, optional): 章节名称，如果为None则获取整个文档内容
            
        Returns:
            str: 章节内容
        """
        return self._get_word_content(chapter_name)
    
    def _get_word_content(self, chapter_name=None):
        """从Word文档中获取内容"""
        if not self.doc:
            return ""
        
        content_parts = []
        in_target_chapter = chapter_name is None
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检查是否是标题
            if paragraph.style.name.startswith('Heading'):
                if chapter_name and chapter_name in text:
                    # 找到目标章节
                    in_target_chapter = True
                    continue
                elif in_target_chapter and paragraph.style.name.startswith('Heading'):
                    # 遇到下一个标题，停止收集内容
                    break
            elif in_target_chapter:
                # 收集章节内容
                content_parts.append(text)
        
        return self._clean_content('\n'.join(content_parts))
    
    def _clean_content(self, content):
        """清理内容文本"""
        if not content:
            return ""
        
        # 移除多余的空白字符
        import re
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'\n\s*\n', '\n\n', content)
        
        # 限制内容长度
        if len(content) > 10000:
            content = content[:10000] + "..."
        
        return content.strip()

    def has_sub_chapters(self, parent_chapter):
        """
        检查指定章节是否有子章节
        
        Args:
            parent_chapter (str): 父章节名称
            
        Returns:
            bool: 是否有子章节
        """
        if not self.doc:
            return False
        
        parent_found = False
        current_level = 0
        
        for paragraph in self.doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                try:
                    level = int(paragraph.style.name.replace('Heading ', ''))
                except ValueError:
                    level = 1
                
                title = paragraph.text.strip()
                if not title:
                    continue
                
                # 精确匹配父章节名称
                if title == parent_chapter:
                    parent_found = True
                    current_level = level
                    continue
                elif parent_found:
                    # 如果已经找到父章节，检查当前章节是否是子章节
                    if level > current_level:
                        return True  # 找到子章节
                    elif level <= current_level:
                        # 遇到同级或更高级的标题，停止搜索
                        break
        
        return False

@app.route('/api/chapters', methods=['GET'])
def get_chapters():
    """
    获取目录章节API
    
    参数:
    - url (必选): Word文件URL
    - chapter (可选): 父级章节名称
    
    返回:
    - JSON格式的章节列表，如果指定章节没有子章节则返回该章节的内容
    """
    try:
        # 获取参数
        file_url = request.args.get('url')
        parent_chapter = request.args.get('chapter')
        
        # 验证必选参数
        if not file_url:
            return jsonify({
                'success': False,
                'error': '缺少必选参数: url'
            }), 400
        
        # 验证URL格式
        try:
            from urllib.parse import urlparse
            parsed_url = urlparse(file_url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return jsonify({
                    'success': False,
                    'error': '无效的URL格式'
                }), 400
        except Exception:
            return jsonify({
                'success': False,
                'error': '无效的URL格式'
            }), 400
        
        # 解析文档
        parser = DocumentParser(file_url)
        
        # 如果指定了章节，先检查是否有子章节
        if parent_chapter:
            # 先检查是否有子章节
            has_sub = parser.has_sub_chapters(parent_chapter)
            
            if has_sub:
                # 有子章节，获取子章节列表
                chapters = parser.get_chapters(parent_chapter, return_content=False)
                return jsonify({
                    'success': True,
                    'data': {
                        'url': file_url,
                        'document_type': 'word',
                        'parent_chapter': parent_chapter,
                        'type': 'chapters',
                        'chapters': chapters,
                        'count': len(chapters),
                        'has_sub_chapters': True
                    }
                })
            else:
                # 没有子章节，获取内容
                content = parser.get_content(parent_chapter)
                return jsonify({
                    'success': True,
                    'data': {
                        'url': file_url,
                        'document_type': 'word',
                        'parent_chapter': parent_chapter,
                        'type': 'content',
                        'is_leaf': True,
                        'content': content,
                        'content_length': len(content),
                        'has_sub_chapters': False
                    }
                })
        else:
            # 获取第一层级章节
            chapters = parser.get_chapters()
            return jsonify({
                'success': True,
                'data': {
                    'url': file_url,
                    'document_type': 'word',
                    'parent_chapter': parent_chapter,
                    'type': 'chapters',
                    'chapters': chapters,
                    'count': len(chapters)
                }
            })
        
    except Exception as e:
        logger.error(f"获取章节失败: {e}")
        return jsonify({
            'success': False,
            'error': f'获取章节失败: {str(e)}'
        }), 500

@app.route('/api/content', methods=['GET'])
def get_content():
    """
    获取章节内容API
    
    参数:
    - url (必选): Word文件URL
    - chapter (可选): 章节名称
    
    返回:
    - JSON格式的章节内容
    """
    try:
        # 获取参数
        file_url = request.args.get('url')
        chapter_name = request.args.get('chapter')
        
        # 验证必选参数
        if not file_url:
            return jsonify({
                'success': False,
                'error': '缺少必选参数: url'
            }), 400
        
        # 验证URL格式
        try:
            from urllib.parse import urlparse
            parsed_url = urlparse(file_url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return jsonify({
                    'success': False,
                    'error': '无效的URL格式'
                }), 400
        except Exception:
            return jsonify({
                'success': False,
                'error': '无效的URL格式'
            }), 400
        
        # 解析文档
        parser = DocumentParser(file_url)
        content = parser.get_content(chapter_name)
        
        return jsonify({
            'success': True,
            'data': {
                'url': file_url,
                'document_type': 'word',
                'chapter': chapter_name,
                'content': content,
                'content_length': len(content)
            }
        })
        
    except Exception as e:
        logger.error(f"获取内容失败: {e}")
        return jsonify({
            'success': False,
            'error': f'获取内容失败: {str(e)}'
        }), 500

@app.route('/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({
        'status': 'healthy',
        'message': '服务运行正常',
        'supported_formats': ['Word (.docx, .doc)']
    })

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'success': False,
        'error': '接口不存在'
    }), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        'success': False,
        'error': '服务器内部错误'
    }), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000) 